from flask import Flask, request, jsonify
from flask_cors import CORS
import os
from dotenv import load_dotenv
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
import re
import sys
from gigachat import GigaChat

# Настройка логирования с правильной кодировкой
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)

load_dotenv()

app = Flask(__name__)
CORS(app)

# Создаем директории для заявок, если их нет
REQUESTS_DIR = "requests"
DOCX_DIR = "docx_files"
for directory in [REQUESTS_DIR, DOCX_DIR]:
    if not os.path.exists(directory):
        os.makedirs(directory)

# Инициализация API ключа
GIGACHAT_API_KEY = os.getenv('GIGACHAT_API_KEY')
if not GIGACHAT_API_KEY:
    raise Exception("Не найден API ключ GigaChat. Пожалуйста, добавьте GIGACHAT_API_KEY в .env файл")

# Путь к сертификату
CA_BUNDLE_FILE = "russian_trusted_root_ca.cer"

# Проверяем наличие сертификата
if not os.path.exists(CA_BUNDLE_FILE):
    raise Exception(f"Сертификат безопасности не найден: {CA_BUNDLE_FILE}")

# Инициализация GigaChat с сертификатом
try:
    giga = GigaChat(
        credentials=GIGACHAT_API_KEY,

        verify_ssl_certs=False
    )
    logging.info("GigaChat успешно инициализирован с сертификатом безопасности")
except Exception as e:
    logging.error(f"Ошибка инициализации GigaChat: {str(e)}")
    raise Exception(f"Не удалось инициализировать GigaChat: {str(e)}")

def validate_input(text):
    """Валидация входного текста"""
    if not text or not isinstance(text, str):
        return False, "Пустой или некорректный текст"
    if len(text.strip()) < 10:
        return False, "Текст слишком короткий"
    return True, ""

def preprocess_text(text):
    """Предварительная обработка текста"""
    # Удаление лишних пробелов
    text = re.sub(r'\s+', ' ', text).strip()
    # Удаление специальных символов
    text = re.sub(r'[^\w\s.,!?-]', '', text)
    return text

def generate_text(prompt, max_length=2000):
    try:
        # Формируем системный промпт
        system_prompt = """# Ты - помощник по классификации обращений пользователей

## Задача
Твоя задача классифицировать обращения пользователей в одну из следующих категорий:
- Создание ПО
- Создание модели машинного обучения
- Создание статистического отчета
- Другое

После определения категории задай вопросы интервьюеру, которые позволят лучше разобраться в теме обращения.

## Инструкция
Для правильного выполнения задания следуй этим шагам:
1. Прочитай обращение внимательно.
2. Найди ключевые слова или фразы, указывающие на тематику обращения.
3. Классифицируй обращение в соответствующую категорию.
4. Задай вопросы, которые помогут глубже понять проблему или запрос пользователя.

## Формат ответа
Ответ должен содержать две части:
1. Категория обращения (одна из предложенных выше).
2. Список вопросов для интервьюера."""

        # Формируем полный запрос
        full_prompt = f"{system_prompt}\n\nЗапрос пользователя: {prompt}"
        
        logging.info("Отправка запроса к GigaChat")
        response = giga.chat(full_prompt)
        
        generated_text = response.choices[0].message.content
        
        # Очищаем текст от возможных артефактов
        generated_text = re.sub(r'\s+', ' ', generated_text)
        generated_text = re.sub(r'[^\w\s.,!?-]', '', generated_text)
        
        return generated_text
        
    except Exception as e:
        logging.error(f"Ошибка генерации текста: {str(e)}")
        raise Exception(f"Ошибка при генерации текста: {str(e)}")

def create_docx(filename, user_request, processed_text, department, start_date, end_date):
    try:
        doc = Document()
        
        # Настройка стилей
        title_style = doc.styles['Title']
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        
        heading_style = doc.styles['Heading 1']
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        
        # Добавление заголовка
        title = doc.add_heading('Заявка на разработку', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавление даты
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_paragraph.add_run(f'Дата создания: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        # Добавление информации об отделе и сроках
        doc.add_heading('Общая информация:', level=1)
        info_paragraph = doc.add_paragraph()
        info_paragraph.add_run(f'Отдел/Департамент: {department}\n')
        info_paragraph.add_run(f'Сроки выполнения: с {start_date} по {end_date}')
        
        # Добавление исходной заявки
        doc.add_heading('Исходная заявка:', level=1)
        doc.add_paragraph(user_request)
        
        # Добавление уточняющих вопросов
        doc.add_heading('Уточняющие вопросы:', level=1)
        doc.add_paragraph(processed_text)
        
        # Сохранение документа
        doc.save(filename)
        print(f"Документ сохранен: {filename}")
        return True
    except Exception as e:
        print(f"Ошибка создания DOCX: {str(e)}")
        return False

@app.route('/api/submit', methods=['POST'])
def submit_request():
    try:
        logging.info("Получен POST запрос")
        data = request.json
        
        if not data or 'request' not in data:
            return jsonify({"success": False, "error": "Отсутствуют данные в запросе"}), 400
            
        user_request = data.get('request')
        department = data.get('department', 'Не указан')
        start_date = data.get('startDate', 'Не указана')
        end_date = data.get('endDate', 'Не указана')
        
        # Валидация входных данных
        is_valid, error_message = validate_input(user_request)
        if not is_valid:
            return jsonify({"success": False, "error": error_message}), 400
            
        # Предварительная обработка текста
        user_request = preprocess_text(user_request)
        
        logging.info(f"Получена заявка: {user_request[:100]}...")
        
        # Обработка запроса с помощью модели
        try:
            processed_text = generate_text(user_request)
        except Exception as e:
            logging.error(f"Ошибка при обработке запроса: {str(e)}")
            return jsonify({"success": False, "error": str(e)}), 500
        
        # Создаем имя файла с текущей датой и временем
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Сохраняем в txt
        txt_filename = f"{REQUESTS_DIR}/request_{timestamp}.txt"
        try:
            with open(txt_filename, 'w', encoding='utf-8') as f:
                f.write(f"Дата и время: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Отдел/Департамент: {department}\n")
                f.write(f"Сроки выполнения: с {start_date} по {end_date}\n")
                f.write(f"Исходная заявка:\n{user_request}\n\n")
                f.write(f"Обработанный ответ:\n{processed_text}\n")
            print(f"TXT файл сохранен: {txt_filename}")
        except Exception as e:
            print(f"Ошибка сохранения TXT: {str(e)}")
            return jsonify({"success": False, "error": f"Ошибка сохранения TXT файла: {str(e)}"}), 500
        
        # Сохраняем в docx
        docx_filename = f"{DOCX_DIR}/ТЗ_{timestamp}.docx"
        if not create_docx(docx_filename, user_request, processed_text, department, start_date, end_date):
            print("Ошибка создания DOCX файла")
            return jsonify({"success": False, "error": "Ошибка создания DOCX файла"}), 500
        
        print("Запрос успешно обработан")
        return jsonify({
            "success": True, 
            "message": "Заявка успешно обработана и сохранена"
        })
    
    except Exception as e:
        print(f"Общая ошибка: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True) 